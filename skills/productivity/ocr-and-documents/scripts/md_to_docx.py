#!/usr/bin/env python3
"""
Markdown → DOCX converter with full table, heading, list, bold, and code block support.
Proven pattern used across multiple sessions for generating Word research reports.

Usage:
    python3 md_to_docx.py input.md output.docx

Requires: python-docx (install: uv pip install --system python-docx)
"""
import sys, re

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx not installed. Run: uv pip install --system python-docx", file=sys.stderr)
    sys.exit(1)


def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Table
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            rows = []
            for tl in table_lines:
                if all(c in '|-: ' for c in tl):
                    continue
                cells = [c.strip() for c in tl.split('|')]
                cells = [c for c in cells if c != '']
                rows.append(cells)

            if rows:
                max_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for ri, row in enumerate(rows):
                    for ci in range(max_cols):
                        cell = table.cell(ri, ci)
                        if ci < len(row):
                            cell_text = row[ci].replace('**', '')
                            cell.text = cell_text
                            if ri == 0:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.bold = True
                                        run.font.size = Pt(10)
                            else:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.size = Pt(10)
                        else:
                            cell.text = ''

                doc.add_paragraph()
            continue

        # Code block
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1

            for cl in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(cl)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(1)
            doc.add_paragraph()
            continue

        # List items
        if stripped.startswith('- '):
            text = stripped[2:].replace('**', '')
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        if stripped.startswith('  - '):
            text = stripped[4:].replace('**', '')
            doc.add_paragraph(text, style='List Bullet 2')
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped).replace('**', '')
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Normal paragraph with bold support
        text = stripped
        p = doc.add_paragraph()
        parts = text.split('**')
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            if idx % 2 == 1:
                run.bold = True
            run.font.size = Pt(11)

        i += 1

    doc.save(docx_path)
    print(f'saved {docx_path}')

    # Verify
    doc2 = Document(docx_path)
    print(f'paragraphs {len(doc2.paragraphs)} tables {len(doc2.tables)}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 md_to_docx.py input.md output.docx")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
